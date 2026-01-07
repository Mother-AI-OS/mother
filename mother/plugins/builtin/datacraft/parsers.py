"""Document parsers for various file formats."""

from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from pypdf import PdfReader


@dataclass
class ParsedDocument:
    """Result of parsing a document."""

    content: str
    pages: int
    tables: list[list[list[str]]]
    metadata: dict[str, Any]
    entities: list[dict[str, str]]
    file_hash: str


def compute_file_hash(file_path: Path) -> str:
    """Compute SHA256 hash of a file."""
    sha256 = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            sha256.update(chunk)
    return sha256.hexdigest()


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> list[str]:
    """Split text into overlapping chunks.

    Args:
        text: Text to chunk
        chunk_size: Target size of each chunk in characters
        overlap: Overlap between chunks

    Returns:
        List of text chunks
    """
    if not text:
        return []

    # Clean up whitespace
    text = re.sub(r"\s+", " ", text).strip()

    if len(text) <= chunk_size:
        return [text]

    chunks = []
    start = 0

    while start < len(text):
        end = start + chunk_size

        # Try to break at sentence boundary
        if end < len(text):
            # Look for sentence end within last 20% of chunk
            search_start = start + int(chunk_size * 0.8)
            search_text = text[search_start:end]

            # Find last sentence boundary
            for sep in [". ", ".\n", "! ", "? ", "\n\n"]:
                pos = search_text.rfind(sep)
                if pos != -1:
                    end = search_start + pos + len(sep)
                    break

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        start = end - overlap

    return chunks


def extract_entities_simple(text: str) -> list[dict[str, str]]:
    """Extract simple entities from text using regex patterns.

    This is a basic implementation. For production, use spaCy or similar.
    """
    entities = []

    # Email addresses
    emails = re.findall(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b", text)
    for email in set(emails):
        entities.append({"type": "EMAIL", "value": email})

    # Phone numbers (various formats)
    phones = re.findall(r"\b(?:\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b", text)
    for phone in set(phones):
        entities.append({"type": "PHONE", "value": phone})

    # Dates (various formats)
    dates = re.findall(
        r"\b\d{1,2}[./-]\d{1,2}[./-]\d{2,4}\b|\b\d{4}[./-]\d{1,2}[./-]\d{1,2}\b",
        text,
    )
    for date in set(dates):
        entities.append({"type": "DATE", "value": date})

    # Money amounts (EUR, USD, etc.)
    amounts = re.findall(
        r"[$\u20AC\u00A3]\s*[\d,]+\.?\d*|\d+[.,]\d{2}\s*(?:EUR|USD|GBP|CHF)",
        text,
        re.IGNORECASE,
    )
    for amount in set(amounts):
        entities.append({"type": "MONEY", "value": amount})

    # IBAN
    ibans = re.findall(r"\b[A-Z]{2}\d{2}[A-Z0-9]{4,30}\b", text)
    for iban in set(ibans):
        if len(iban) >= 15:  # IBANs are typically 15-34 chars
            entities.append({"type": "IBAN", "value": iban})

    return entities


def detect_document_type(filename: str, content: str) -> str:
    """Detect document type from filename and content.

    Returns:
        Document type: invoice, receipt, bank_statement, contract, letter, report, other
    """
    filename_lower = filename.lower()
    content_lower = content.lower()[:2000]

    # Check filename patterns
    if any(x in filename_lower for x in ["invoice", "rechnung", "inv_", "inv-"]):
        return "invoice"
    if any(x in filename_lower for x in ["receipt", "quittung", "beleg"]):
        return "receipt"
    if any(x in filename_lower for x in ["statement", "kontoauszug", "bank"]):
        return "bank_statement"
    if any(x in filename_lower for x in ["contract", "vertrag", "agreement"]):
        return "contract"

    # Check content patterns
    invoice_keywords = ["invoice", "rechnung", "total amount", "gesamtbetrag", "due date", "fällig"]
    if sum(1 for kw in invoice_keywords if kw in content_lower) >= 2:
        return "invoice"

    receipt_keywords = ["receipt", "quittung", "paid", "bezahlt", "thank you for your purchase"]
    if sum(1 for kw in receipt_keywords if kw in content_lower) >= 2:
        return "receipt"

    bank_keywords = ["account balance", "kontostand", "transaction", "überweisung", "saldo"]
    if sum(1 for kw in bank_keywords if kw in content_lower) >= 2:
        return "bank_statement"

    contract_keywords = ["hereby agree", "vereinbaren", "terms and conditions", "party", "vertragspartner"]
    if sum(1 for kw in contract_keywords if kw in content_lower) >= 2:
        return "contract"

    return "other"


def parse_pdf(file_path: Path) -> ParsedDocument:
    """Parse a PDF document.

    Args:
        file_path: Path to PDF file

    Returns:
        Parsed document with content, metadata, and entities
    """
    reader = PdfReader(str(file_path))

    # Extract text from all pages
    pages_text = []
    for page in reader.pages:
        text = page.extract_text() or ""
        pages_text.append(text)

    full_text = "\n\n".join(pages_text)

    # Extract metadata
    metadata = {}
    if reader.metadata:
        if reader.metadata.title:
            metadata["title"] = reader.metadata.title
        if reader.metadata.author:
            metadata["author"] = reader.metadata.author
        if reader.metadata.creator:
            metadata["creator"] = reader.metadata.creator
        if reader.metadata.creation_date:
            metadata["created"] = str(reader.metadata.creation_date)

    # Extract tables (basic - look for tabular patterns)
    tables = _extract_tables_from_text(full_text)

    # Extract entities
    entities = extract_entities_simple(full_text)

    return ParsedDocument(
        content=full_text,
        pages=len(reader.pages),
        tables=tables,
        metadata=metadata,
        entities=entities,
        file_hash=compute_file_hash(file_path),
    )


def parse_docx(file_path: Path) -> ParsedDocument:
    """Parse a DOCX document.

    Args:
        file_path: Path to DOCX file

    Returns:
        Parsed document with content, metadata, and entities
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")

    doc = Document(str(file_path))

    # Extract paragraphs
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n\n".join(paragraphs)

    # Extract tables
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        if table_data:
            tables.append(table_data)

    # Extract metadata
    metadata = {}
    if doc.core_properties:
        if doc.core_properties.title:
            metadata["title"] = doc.core_properties.title
        if doc.core_properties.author:
            metadata["author"] = doc.core_properties.author
        if doc.core_properties.created:
            metadata["created"] = str(doc.core_properties.created)

    # Extract entities
    entities = extract_entities_simple(full_text)

    return ParsedDocument(
        content=full_text,
        pages=len(doc.sections),  # Approximate page count
        tables=tables,
        metadata=metadata,
        entities=entities,
        file_hash=compute_file_hash(file_path),
    )


def parse_text(file_path: Path) -> ParsedDocument:
    """Parse a plain text file.

    Args:
        file_path: Path to text file

    Returns:
        Parsed document with content and entities
    """
    content = file_path.read_text(encoding="utf-8", errors="replace")
    entities = extract_entities_simple(content)

    return ParsedDocument(
        content=content,
        pages=1,
        tables=[],
        metadata={"encoding": "utf-8"},
        entities=entities,
        file_hash=compute_file_hash(file_path),
    )


def parse_document(file_path: Path) -> ParsedDocument:
    """Parse a document based on its extension.

    Args:
        file_path: Path to document

    Returns:
        Parsed document

    Raises:
        ValueError: If file type is not supported
    """
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return parse_pdf(file_path)
    elif suffix in (".docx", ".doc"):
        return parse_docx(file_path)
    elif suffix in (".txt", ".md", ".rst", ".csv"):
        return parse_text(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def _extract_tables_from_text(text: str) -> list[list[list[str]]]:
    """Extract table-like structures from text.

    This is a simple heuristic that looks for lines with consistent delimiters.
    """
    tables = []
    current_table = []

    for line in text.split("\n"):
        # Check if line looks like a table row
        # Multiple cells separated by | or tabs or multiple spaces
        cells = None

        if "|" in line:
            cells = [c.strip() for c in line.split("|") if c.strip()]
        elif "\t" in line:
            cells = [c.strip() for c in line.split("\t") if c.strip()]
        elif "  " in line:
            # Multiple spaces might indicate columns
            cells = [c.strip() for c in re.split(r"\s{2,}", line) if c.strip()]

        if cells and len(cells) >= 2:
            current_table.append(cells)
        else:
            # End of table
            if len(current_table) >= 2:  # At least header + 1 row
                tables.append(current_table)
            current_table = []

    # Don't forget last table
    if len(current_table) >= 2:
        tables.append(current_table)

    return tables


# Supported file extensions
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md", ".rst", ".csv"}


def is_supported(file_path: Path) -> bool:
    """Check if a file type is supported."""
    return file_path.suffix.lower() in SUPPORTED_EXTENSIONS
